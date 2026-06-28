#!/usr/bin/env python3
"""Add headers and page numbers to thesis DOCX — robust version."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INPUT = r"E:\数据\CampusMicroMart\毕业论文.docx"
OUTPUT = r"E:\数据\CampusMicroMart\毕业论文.docx"
TITLE = "基于Spring Cloud Alibaba的微服务化校园二手交易平台设计与实现"

def Ox(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    return el


def main():
    doc = Document(INPUT)
    body = doc.element.body
    
    # Find key paragraph indices
    idx_abstract = idx_toc = idx_body = -1
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == '摘  要' and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            idx_abstract = i
        if t == '目  录' and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            idx_toc = i
        if '1 引言' in t and p.alignment == WD_ALIGN_PARAGRAPH.CENTER and i > idx_toc:
            idx_body = i
            break
    
    print(f"Abstract at {idx_abstract}, TOC at {idx_toc}, Body at {idx_body}")
    
    # ── Insert section break #1: between cover and abstract ──
    # Find the empty paragraph just before "摘  要"
    break1_idx = idx_abstract - 1
    para1 = doc.paragraphs[break1_idx]
    pPr1 = para1._p.get_or_add_pPr()
    sectPr1 = Ox('w:sectPr')
    sectPr1.append(Ox('w:type', **{'w:val': 'nextPage'}))
    sectPr1.append(Ox('w:pgSz', **{'w:w': '11906', 'w:h': '16838'}))
    pPr1.append(sectPr1)
    
    # ── Insert section break #2: between TOC and body ──
    # Find the empty paragraph just before "1 引言"
    break2_idx = idx_body - 1
    para2 = doc.paragraphs[break2_idx]
    pPr2 = para2._p.get_or_add_pPr()
    sectPr2 = Ox('w:sectPr')
    sectPr2.append(Ox('w:type', **{'w:val': 'nextPage'}))
    sectPr2.append(Ox('w:pgSz', **{'w:w': '11906', 'w:h': '16838'}))
    pPr2.append(sectPr2)
    
    # Save to capture section breaks, then reload
    doc.save(OUTPUT)
    
    # Remove body-level sectPr so all sections come from paragraph-level breaks
    # Need to do this BEFORE re-opening
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    doc2 = Document(OUTPUT)
    body2 = doc2.element.body
    for sp in body2.findall(qn('w:sectPr')):
        body2.remove(sp)
    doc2.save(OUTPUT)
    
    doc = Document(OUTPUT)
    
    sections = doc.sections
    print(f"Sections after reload: {len(sections)}")
    
    # ── Section 0: Cover ──
    s0 = sections[0]
    s0.different_first_page_header_footer = True
    # Empty first page header
    for p in s0.first_page_header.paragraphs:
        p.clear()
    s0.first_page_header.is_linked_to_previous = False
    # Empty header
    for p in s0.header.paragraphs:
        p.clear()
    s0.header.is_linked_to_previous = False
    # Empty footers
    for p in s0.first_page_footer.paragraphs:
        p.clear()
    s0.first_page_footer.is_linked_to_previous = False
    for p in s0.footer.paragraphs:
        p.clear()
    s0.footer.is_linked_to_previous = False
    
    # ── Section 1: Abstract + TOC ──
    if len(sections) >= 2:
        s1 = sections[1]
        s1.different_first_page_header_footer = False
        # Header: empty
        for p in s1.header.paragraphs:
            p.clear()
        s1.header.is_linked_to_previous = False
        # Footer: Roman-style page numbers
        s1.footer.is_linked_to_previous = False
        for p in s1.footer.paragraphs:
            p.clear()
        fp = s1.footer.paragraphs[0] if s1.footer.paragraphs else s1.footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(); r.font.size = Pt(9)
        rPr = r._r.get_or_add_rPr()
        rf = Ox('w:rFonts', **{'w:eastAsia': '宋体', 'w:ascii': 'Times New Roman', 'w:hAnsi': 'Times New Roman'})
        rPr.insert(0, rf)
        # PAGE field
        r2 = fp.add_run()
        fc1 = Ox('w:fldChar', **{'w:fldCharType': 'begin'})
        r2._r.append(fc1)
        r3 = fp.add_run()
        it = Ox('w:instrText', **{'xml:space': 'preserve'})
        it.text = ' PAGE \\* ROMAN '
        r3._r.append(it)
        r4 = fp.add_run()
        fc2 = Ox('w:fldChar', **{'w:fldCharType': 'end'})
        r4._r.append(fc2)
    
    # ── Section 2: Body + References ──
    if len(sections) >= 3:
        s2 = sections[2]
        s2.different_first_page_header_footer = False
        # Header: thesis title + bottom border
        s2.header.is_linked_to_previous = False
        for p in s2.header.paragraphs:
            p.clear()
        hp = s2.header.paragraphs[0] if s2.header.paragraphs else s2.header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(TITLE)
        hr.font.size = Pt(9); hr.font.name = 'Times New Roman'
        hrPr = hr._r.get_or_add_rPr()
        hrf = Ox('w:rFonts', **{'w:eastAsia': '宋体', 'w:ascii': 'Times New Roman', 'w:hAnsi': 'Times New Roman'})
        hrPr.insert(0, hrf)
        # Bottom border
        hpPr = hp._p.get_or_add_pPr()
        pBdr = Ox('w:pBdr')
        pBdr.append(Ox('w:bottom', **{'w:val': 'single', 'w:sz': '6', 'w:space': '4', 'w:color': '000000'}))
        hpPr.append(pBdr)
        
        # Footer: centered Arabic page number
        s2.footer.is_linked_to_previous = False
        for p in s2.footer.paragraphs:
            p.clear()
        fp2 = s2.footer.paragraphs[0] if s2.footer.paragraphs else s2.footer.add_paragraph()
        fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp2.add_run(); r.font.size = Pt(9)
        rPr2 = r._r.get_or_add_rPr()
        rf2 = Ox('w:rFonts', **{'w:eastAsia': '宋体', 'w:ascii': 'Times New Roman', 'w:hAnsi': 'Times New Roman'})
        rPr2.insert(0, rf2)
        # PAGE field
        r2 = fp2.add_run()
        r2._r.append(Ox('w:fldChar', **{'w:fldCharType': 'begin'}))
        r3 = fp2.add_run()
        it = Ox('w:instrText', **{'xml:space': 'preserve'})
        it.text = ' PAGE '
        r3._r.append(it)
        r4 = fp2.add_run()
        r4._r.append(Ox('w:fldChar', **{'w:fldCharType': 'end'}))
    
    # Ensure page size for all sections
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)
    
    doc.save(OUTPUT)
    print(f"✅ Saved: {OUTPUT}")

if __name__ == '__main__':
    main()
