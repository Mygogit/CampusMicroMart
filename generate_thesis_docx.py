#!/usr/bin/env python3
"""Generate thesis DOCX per 内蒙古财经大学 formatting standards."""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════ Configuration: 内蒙古财经大学 Standards ═══════════════
MARKDOWN_FILE = r"E:\数据\CampusMicroMart\毕业论文.md"
OUTPUT_FILE = r"E:\数据\CampusMicroMart\毕业论文.docx"
THESIS_TITLE = "基于Spring Cloud Alibaba的微服务化校园二手交易平台设计与实现"

# ── Fonts ──
FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_KAI = "楷体"
FONT_FANGSONG = "仿宋"
FONT_EN = "Times New Roman"
FONT_CODE = "Consolas"

# ── Sizes (per IMUFE standard) ──
SZ_COVER_INSTITUTION = Pt(15)    # 小三号 宋体 "内蒙古财经大学教务处制表"
SZ_COVER_TITLE_CN = Pt(16)       # 三号 宋体加粗 "本科毕业论文（设计）"
SZ_COVER_INFO_LABEL = Pt(16)     # 三号 黑体加粗 个人信息标签
SZ_COVER_INFO_TEXT = Pt(16)      # 三号 楷体 个人信息内容
SZ_TITLE = Pt(22)                # 二号 黑体加粗 封面标题
SZ_SUBTITLE = Pt(18)             # 小二号 黑体 副标题

SZ_ABSTRACT_HEADING = Pt(16)     # 三号 黑体 "摘  要"
SZ_ABSTRACT_BODY = Pt(12)        # 小四号 宋体
SZ_KEYWORD_LABEL = Pt(12)        # 小四号 黑体加粗
SZ_KEYWORD_TEXT = Pt(12)         # 小四号 宋体加粗

SZ_EN_HEADING = Pt(16)           # 三号 Times New Roman bold
SZ_EN_BODY = Pt(12)              # 小四号 Times New Roman
SZ_EN_KEY_LABEL = Pt(12)         # 小四号 黑体加粗
SZ_EN_KEY_TEXT = Pt(12)          # 小四号 Times New Roman 斜体

SZ_TOC_HEADING = Pt(16)          # 三号 黑体
SZ_TOC_BODY = Pt(12)             # 小四号 仿宋

SZ_BODY_TITLE = Pt(18)           # 小二号 黑体 正文标题（另起页）
SZ_BODY_SUBTITLE = Pt(16)        # 三号 黑体
SZ_BODY = Pt(12)                 # 小四号 宋体 正文
SZ_H1 = Pt(15)                   # 小三号 黑体 一级标题
SZ_H2 = Pt(14)                   # 四号 黑体 二级标题
SZ_H3 = Pt(12)                   # 小四号 宋体加粗 三级标题

SZ_REF_HEADING = Pt(14)          # 四号 黑体
SZ_REF_BODY = Pt(12)             # 小四号 仿宋（中文）/ Times New Roman（英文）
SZ_REF_EN = Pt(12)               # 小四号 Times New Roman

SZ_THANKS_HEADING = Pt(14)       # 四号 黑体 居中
SZ_THANKS_BODY = Pt(12)          # 小四号 楷体

SZ_TABLE_CAPTION = Pt(10.5)      # 五号 黑体
SZ_TABLE_BODY = Pt(10.5)         # 五号 宋体

SZ_FIGURE_CAPTION = Pt(10.5)      # 五号 黑体

SZ_CODE = Pt(9)                   # 小五号 Consolas

SZ_HEADER_FOOTER = Pt(10.5)      # 五号 黑体 页眉

SZ_FOOTNOTE = Pt(9)              # 小五号 宋体

# ── Spacing ──
LINE_SPACING_BODY = 1.5
FIRST_LINE_INDENT = Cm(0.74)     # 2 characters at 小四号

# ── Margins ──
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)
HEADER_DISTANCE = Cm(2.0)
FOOTER_DISTANCE = Cm(1.75)


# ═══════════════ Helpers ═══════════════

def Ox(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    return el

def set_run_font(run, cn=None, en=None, size=None, bold=False, italic=False, color=None):
    if size is not None: run.font.size = size
    if bold: run.bold = True
    if italic: run.italic = True
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    if cn: rFonts.set(qn('w:eastAsia'), cn)
    if en: rFonts.set(qn('w:ascii'), en); rFonts.set(qn('w:hAnsi'), en)
    if color: run.font.color.rgb = color

def set_para_sp(para, ls=LINE_SPACING_BODY, before=0, after=0):
    pf = para.paragraph_format
    pf.line_spacing = ls
    pf.space_before = before
    pf.space_after = after

def add_heading(doc, text, level):
    """Add heading per IMUFE standard."""
    para = doc.add_paragraph()
    sizes = {1: (SZ_H1, FONT_HEI), 2: (SZ_H2, FONT_HEI), 3: (SZ_H3, FONT_SONG)}
    sz, font = sizes.get(level, (SZ_H3, FONT_SONG))
    if level <= 2: para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    elif level == 3: para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    r = para.add_run(text)
    set_run_font(r, cn=font, en=FONT_EN, size=sz, bold=(level <= 2))
    set_para_sp(para, 1.25, before=Pt(6), after=Pt(2))

def add_body(doc, text, indent=True, font_cn=FONT_SONG, font_en=FONT_EN, size=SZ_BODY):
    """Add body paragraph."""
    para = doc.add_paragraph()
    set_para_sp(para, LINE_SPACING_BODY, before=0, after=0)
    if indent:
        para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            set_run_font(r, cn=font_cn, en=font_en, size=size, bold=True)
        else:
            r = para.add_run(part)
            set_run_font(r, cn=font_cn, en=font_en, size=size)
    return para

def add_centered_heading(doc, text, size, font_cn, bold=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(text)
    set_run_font(r, cn=font_cn, en=FONT_EN, size=size, bold=bold)
    set_para_sp(para, LINE_SPACING_BODY, before=Pt(12), after=Pt(6))

def add_ref_heading(doc, text):
    """参考文献标题：四号黑体，顶格排印（左对齐），不加冒号"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = para.add_run(text)
    set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_REF_HEADING, bold=True)
    set_para_sp(para, LINE_SPACING_BODY, before=Pt(0), after=Pt(6))

def add_code(doc, text):
    para = doc.add_paragraph()
    set_para_sp(para, 1.0, before=Pt(2), after=Pt(2))
    para.paragraph_format.left_indent = FIRST_LINE_INDENT
    r = para.add_run(text)
    set_run_font(r, cn=FONT_CODE, en=FONT_CODE, size=SZ_CODE, color=RGBColor(0x55, 0x55, 0x55))
    # Gray background
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def add_ref(doc, text):
    para = doc.add_paragraph()
    set_para_sp(para, LINE_SPACING_BODY, before=0, after=0)
    r = para.add_run(text)
    set_run_font(r, cn=FONT_FANGSONG, en=FONT_EN, size=SZ_REF_BODY)

def add_toc_item(doc, text, level=1):
    para = doc.add_paragraph()
    set_para_sp(para, LINE_SPACING_BODY, before=0, after=0)
    if level == 2: para.paragraph_format.left_indent = FIRST_LINE_INDENT
    r = para.add_run(text)
    set_run_font(r, cn=FONT_FANGSONG, en=FONT_EN, size=SZ_TOC_BODY)

def add_kw_cn(doc, text):
    para = doc.add_paragraph()
    set_para_sp(para, LINE_SPACING_BODY, before=Pt(4), after=Pt(8))
    r = para.add_run('关键词：')
    set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_KEYWORD_LABEL, bold=True)
    r2 = para.add_run(text)
    set_run_font(r2, cn=FONT_SONG, en=FONT_EN, size=SZ_KEYWORD_TEXT, bold=True)

def add_kw_en(doc, text):
    para = doc.add_paragraph()
    set_para_sp(para, LINE_SPACING_BODY, before=Pt(4), after=Pt(8))
    r = para.add_run('Key Words: ')
    set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_EN_KEY_LABEL, bold=True)
    r2 = para.add_run(text)
    set_run_font(r2, cn=FONT_EN, en=FONT_EN, size=SZ_EN_KEY_TEXT, italic=True)

def add_table_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_sp(para, 1.0, before=Pt(8), after=Pt(2))
    r = para.add_run(text)
    set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_TABLE_CAPTION, bold=True)

def add_figure_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_sp(para, 1.0, before=Pt(2), after=Pt(6))
    r = para.add_run(text)
    set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_FIGURE_CAPTION)

def make_table(doc, rows):
    if not rows or len(rows) < 1: return
    cols = max(r.count('|') - 1 for r in rows)
    cols = max(cols, 1)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Style as three-line table
    tbl.style = 'Table Grid'
    for ri, row_text in enumerate(rows):
        cells = [c.strip() for c in row_text.strip().strip('|').split('|')]
        for ci in range(min(len(cells), cols)):
            cell = tbl.cell(ri, ci); cell.text = ''
            p = cell.paragraphs[0]
            set_para_sp(p, 1.0, before=Pt(1), after=Pt(1))
            r = p.add_run(cells[ci])
            if ri == 0:
                set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_TABLE_CAPTION, bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_run_font(r, cn=FONT_SONG, en=FONT_EN, size=SZ_TABLE_BODY)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ═══════════════ Section Setup ═══════════════

def section_cover(section):
    """Cover: no page number, no header."""
    section.different_first_page_header_footer = True
    for h in [section.header, section.first_page_header]:
        h.is_linked_to_previous = False
        for p in h.paragraphs: p.clear()
    for f in [section.footer, section.first_page_footer]:
        f.is_linked_to_previous = False
        for p in f.paragraphs: p.clear()

def section_frontmatter(section):
    """Abstract+TOC: no page number, no header."""
    for h in [section.header]:
        h.is_linked_to_previous = False
        for p in h.paragraphs: p.clear()
    for f in [section.footer]:
        f.is_linked_to_previous = False
        for p in f.paragraphs: p.clear()

def section_body(section):
    """Body: odd=大学名 header, even=论文题 header. Page number odd右下 even左下."""
    section.different_first_page_header_footer = False
    
    # ── Even page header (left-aligned) = thesis title ──
    even_header = section.header
    even_header.is_linked_to_previous = False
    for p in even_header.paragraphs: p.clear()
    hp = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run(THESIS_TITLE)
    set_run_font(hr, cn=FONT_HEI, en=FONT_EN, size=SZ_HEADER_FOOTER)
    # Even header bottom border
    hpPr = hp._p.get_or_add_pPr()
    pBdr = Ox('w:pBdr')
    pBdr.append(Ox('w:bottom', **{'w:val': 'single', 'w:sz': '6', 'w:space': '4', 'w:color': '000000'}))
    hpPr.append(pBdr)
    
    # ── Odd page footer: page number right-aligned ──
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs: p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # PAGE field
    r1 = fp.add_run(); r1._r.append(Ox('w:fldChar', **{'w:fldCharType': 'begin'}))
    r2 = fp.add_run()
    it = Ox('w:instrText', **{'xml:space': 'preserve'}); it.text = ' PAGE '
    r2._r.append(it)
    r3 = fp.add_run(); r3._r.append(Ox('w:fldChar', **{'w:fldCharType': 'end'}))
    for run in fp.runs:
        run.font.size = Pt(9)
        rPr = run._r.get_or_add_rPr()
        rfs = Ox('w:rFonts', **{'w:eastAsia': FONT_SONG, 'w:ascii': FONT_EN, 'w:hAnsi': FONT_EN})
        rPr.insert(0, rfs)

def setup_section(section):
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP; section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT; section.right_margin = MARGIN_RIGHT


# ═══════════════ Cover Page ═══════════════

def build_cover(doc):
    # University name / logo area (placeholder)
    for _ in range(3):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_sp(p, LINE_SPACING_BODY)
    
    # "本科毕业论文（设计）"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_sp(p, LINE_SPACING_BODY, before=Pt(24), after=Pt(24))
    r = p.add_run('本科毕业论文（设计）')
    set_run_font(r, cn=FONT_SONG, en=FONT_EN, size=SZ_COVER_TITLE_CN, bold=True)
    
    # Blank line before title
    p = doc.add_paragraph(); set_para_sp(p, LINE_SPACING_BODY)
    
    # Title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_sp(p, LINE_SPACING_BODY)
    r = p.add_run(THESIS_TITLE)
    set_run_font(r, cn=FONT_HEI, en=FONT_EN, size=SZ_TITLE, bold=True)
    
    # Blank line after title
    p = doc.add_paragraph(); set_para_sp(p, LINE_SPACING_BODY)
    
    # Author info
    info_fields = [
        ('姓    名：', '_____________'),
        ('学    号：', '_____________'),
        ('学    院：', '_____________'),
        ('专    业：', '_____________'),
        ('学    位：', '_____________'),
        ('班    级：', '_____________'),
        ('指导教师：', '_____________'),
        ('导师职称：', '_____________'),
    ]
    for label, value in info_fields:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_sp(p, LINE_SPACING_BODY)
        r1 = p.add_run(label)
        set_run_font(r1, cn=FONT_HEI, en=FONT_EN, size=SZ_COVER_INFO_LABEL, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, cn=FONT_KAI, en=FONT_EN, size=SZ_COVER_INFO_TEXT)
    
    # Institution line
    for _ in range(2):
        p = doc.add_paragraph(); set_para_sp(p, LINE_SPACING_BODY)
    
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_sp(p, LINE_SPACING_BODY)
    r = p.add_run('内蒙古财经大学教务处制表')
    set_run_font(r, cn=FONT_SONG, en=FONT_EN, size=SZ_COVER_INSTITUTION)


# ═══════════════ Number Conversion (Arabic → Chinese) ═══════════════

CN_NUMS = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
           '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']

def to_cn(n):
    """Convert integer 1-20 to Chinese numeral."""
    if 0 <= n < len(CN_NUMS):
        return CN_NUMS[n]
    return str(n)

def convert_chapter_number(text):
    """Convert '1 引言' → '一、引言'"""
    m = re.match(r'^(\d+)\s+(.+)', text)
    if m:
        n = int(m.group(1))
        return f'{to_cn(n)}、{m.group(2)}'
    return text

def convert_section_number(text):
    """Convert '1.1 背景' → '（一）背景', preserving only the subsection number."""
    m = re.match(r'^\d+\.(\d+)\s+(.+)', text)
    if m:
        n = int(m.group(1))
        return f'（{to_cn(n)}）{m.group(2)}'
    return text


# ═══════════════ Main ═══════════════

def generate():
    with open(MARKDOWN_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_EN; style.font.size = SZ_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
    
    # ── Section 0: Cover ──
    build_cover(doc)
    section_cover(doc.sections[0]); setup_section(doc.sections[0])
    
    # Parse frontmatter
    i = 0
    while i < len(lines) and not lines[i].strip().startswith('## 摘要'):
        i += 1
    i += 1
    
    abs_cn = []; kw_cn = ''
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('## Abstract'): i += 1; break
        if line.startswith('**关键词') or line.startswith('**关键字'): kw_cn = line; i += 1; continue
        if line.strip(): abs_cn.append(line.strip())
        i += 1
    
    abs_en = []; kw_en = ''
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('## 目录') or line == '## 目录': i += 1; break
        if line.startswith('**Keywords'): kw_en = line; i += 1; continue
        if line.strip(): abs_en.append(line.strip())
        i += 1
    
    tocs = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('## ') or line.startswith('# '): break
        if line.strip(): tocs.append(line.strip())
        i += 1
    
    # ── Section 1: Chinese Abstract (另起页) ──
    doc.add_section(); section_frontmatter(doc.sections[-1]); setup_section(doc.sections[-1])
    add_centered_heading(doc, '摘  要', SZ_ABSTRACT_HEADING, FONT_HEI)
    for a in abs_cn: add_body(doc, a, indent=True, font_cn=FONT_SONG, size=SZ_ABSTRACT_BODY)
    if kw_cn:
        kw = re.sub(r'\*\*关键词\*\*[：:]?\s*', '', kw_cn).strip()
        add_kw_cn(doc, kw)
    
    # ── Section 2: English Abstract (另起页) ──
    doc.add_section(); section_frontmatter(doc.sections[-1]); setup_section(doc.sections[-1])
    add_centered_heading(doc, 'Abstract', SZ_EN_HEADING, FONT_EN)
    for a in abs_en: add_body(doc, a, indent=True, font_cn=FONT_EN, font_en=FONT_EN, size=SZ_EN_BODY)
    if kw_en:
        kw = re.sub(r'\*\*Keywords\*\*[：:]?\s*', '', kw_en).strip()
        add_kw_en(doc, kw)
    
    # ── Section 3: TOC (单设一页) ──
    doc.add_section(); section_frontmatter(doc.sections[-1]); setup_section(doc.sections[-1])
    add_centered_heading(doc, '目  录', SZ_TOC_HEADING, FONT_HEI)
    # Parse TOC items with level detection
    for t in tocs:
        t_clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t).strip()
        level = 2 if t_clean.startswith('-') or t_clean.startswith('  ') else 1
        t_clean = re.sub(r'^[\-\s]+', '', t_clean)
        if t_clean:
            add_toc_item(doc, t_clean, level)
    
    # ── Section 4: Body + References (从正文开始编页码) ──
    doc.add_section(); section_body(doc.sections[-1]); setup_section(doc.sections[-1])
    
    # Skip the H1 title line from MD (line 1 is `# 基于Spring Cloud...`)
    # In the MD, the title is the first line. We add it manually here.
    add_centered_heading(doc, THESIS_TITLE, SZ_BODY_TITLE, FONT_HEI)
    
    # Process remaining body content
    in_code = False; code_buf = []
    in_table = False; table_buf = []
    in_refs = False
    
    # Find where i is currently (at first ## heading after TOC — should be `## 1 引言`)
    while i < len(lines):
        line = lines[i].rstrip()
        
        if line.startswith('## ') and ('参考文献' in line or '参考' in line):
            in_refs = True
            # Two blank lines before references heading
            p = doc.add_paragraph(); set_para_sp(p, LINE_SPACING_BODY)
            p = doc.add_paragraph(); set_para_sp(p, LINE_SPACING_BODY)
            # 参考文献：四号黑体，顶格排印（不加冒号，不居中）
            add_ref_heading(doc, '参考文献')
            i += 1; continue
        
        if in_refs:
            if line.strip().startswith('['):
                add_ref(doc, line.strip())
            elif line.strip():
                add_ref(doc, line.strip())
            i += 1; continue
        
        # Skip the original H1 title (redundant, already added)
        if line.startswith('# ') and THESIS_TITLE[:8] in line:
            i += 1; continue
        
        # Code fences
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    if cl.strip(): add_code(doc, cl)
                code_buf = []
            i += 1; continue
        if in_code: code_buf.append(line); i += 1; continue
        
        # Tables
        if line.strip().startswith('|') and '---' not in line:
            if not in_table: in_table = True; table_buf = []
            table_buf.append(line.strip()); i += 1; continue
        if line.strip().startswith('|') and '---' in line:
            if in_table: table_buf.append(line.strip())
            i += 1; continue
        if in_table and not line.strip().startswith('|'):
            data_rows = [r for r in table_buf if not re.match(r'^\|[\s\-:|]+\|$', r)]
            if data_rows: make_table(doc, data_rows)
            in_table = False; table_buf = []
        
        # Headings — remap to IMUFE numbering style
        h2 = re.match(r'^## (.+)$', line)
        h3 = re.match(r'^### (.+)$', line)
        h4 = re.match(r'^#### (.+)$', line)
        
        if h2:
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', h2.group(1))
            txt = convert_chapter_number(txt)  # "1 引言" → "一、引言"
            add_heading(doc, txt, 1)
        elif h3:
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', h3.group(1))
            txt = convert_section_number(txt)  # "1.1 背景" → "（一）背景"
            add_heading(doc, txt, 2)
        elif h4:
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', h4.group(1))
            add_heading(doc, txt, 3)
        elif line.strip():
            ls = line.strip()
            if ls in ('---', '***', '___'): pass
            elif ls.startswith('!['): pass
            elif '此处省略' in ls: pass
            elif re.match(r'^\[.*\]\(.*\)$', ls): pass
            else:
                ls = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', ls)
                add_body(doc, ls, indent=True, font_cn=FONT_SONG, size=SZ_BODY)
        i += 1
    
    # ── Section 5: 致谢 ──
    doc.add_section(); section_body(doc.sections[-1]); setup_section(doc.sections[-1])
    add_centered_heading(doc, '致  谢', SZ_THANKS_HEADING, FONT_HEI)
    thanks_text = (
        "在本论文完成之际，谨向所有在课题研究和论文撰写过程中给予我帮助和支持的人表示诚挚的感谢。"
        "特别感谢指导教师在整个毕业设计过程中的悉心指导，从选题、系统设计到论文撰写，导师都给予了宝贵的建议和耐心的帮助。"
        "感谢各位任课老师在大学四年中的教导，为我打下了扎实的专业基础。"
        "感谢同学和朋友们的陪伴与鼓励，在我遇到困难时给予我支持和帮助。"
        "最后，感谢我的家人，你们的理解和支持是我完成学业的最大动力。"
    )
    add_body(doc, thanks_text, indent=True, font_cn=FONT_KAI, size=SZ_THANKS_BODY)
    
    doc.save(OUTPUT_FILE)
    print(f"✅ Generated per IMUFE standards: {OUTPUT_FILE}")
    print(f"   Sections: {len(doc.sections)}")

if __name__ == '__main__':
    generate()
